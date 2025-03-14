import os
import docx
import json
import openpyxl
import argparse
from util import customUtil,formatters
from datetime import datetime
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

CONFIG_DIR = os.path.join("config")

def list_config_files(config_dir="config"):
    """ Lists all available JSON config files and extracts product names. """
    if not os.path.exists(config_dir):
        print(f"❌ Config directory '{config_dir}' not found.")
        return []

    config_files = [f for f in os.listdir(config_dir) if f.endswith(".json")]
    config_info = []

    for file in config_files:
        file_path = os.path.join(config_dir, file)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                config_data = json.load(f)
                product_name = config_data.get("productName", "Unknown Product")
                shouldIgnore = config_data.get("shouldIgnore", False)
                if not shouldIgnore:
                    config_info.append((file, product_name))
        except Exception as e:
            print(f"⚠️  Skipping '{file}': Unable to read or invalid JSON format. ({e})")

    if not config_info:
        print(f"❌ No valid JSON config files found in '{config_dir}'.")
    
    return config_info

def get_user_selected_config(config_info, config_dir="config"):
    """ Prompts the user to select a config file from available options. """
    while True:
        print("\n📌 Available Configurations:")
        for idx, (file, product_name) in enumerate(config_info, start=1):
            file_path = os.path.join(config_dir, file)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    config_data = json.load(f)
                    localInputPath = config_data.get("inputPath", "Unknown Product")
            except Exception as e:
                print(f"⚠️  Skipping '{file}': Unable to read or invalid JSON format. ({e})")
            print(f"{idx}. {product_name} ({localInputPath})")

        choice = input("\nEnter the number of the configuration to use (or 'q' to quit): ")
        if choice.lower() == 'q':
            print("\n🛑 Exiting selection...")
            return None
        try:
            choice = int(choice)
            if 1 <= choice <= len(config_info):
                return os.path.join(config_dir, config_info[choice - 1][0])
            else:
                print("❌ Invalid selection. Please enter a valid number.")
        except ValueError:
            print("❌ Invalid input. Please enter a number.")

def read_data(file_path,config):
    """ Reads input data from a CSV or XLSX file and converts it to a dictionary. """
    if not os.path.exists(file_path):
        return {}, [f"❌ Data file '{file_path}' not found."]
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext in [".xls", ".xlsx"]:
        return read_xlsx(file_path,config)
    else:
        return {}, ["❌ Unsupported file format. Please use XLSX."]

def format_cell_value(cell,number_format):
    if isinstance(cell, datetime):
        return cell.strftime("%d %B %Y")  # Matches "10 October 2024"
    
    if isinstance(cell, (int, float)) and "%" in number_format:
        return f"{cell * 100:.0f}%"  # Convert 0.98 → 98%
    
    if isinstance(cell, (int, float)) and "," in number_format:  # Checks if formatting includes commas
        return  f"{cell:,}" 
    return str(cell).strip() if cell is not None else ""

def read_xlsx(file_path, config):
    """Reads only the sheets mentioned in the config file while preserving formats."""
    wb = openpyxl.load_workbook(file_path, data_only=True)
    data = {}
    errors = []

    required_sheets = set()
    for mapping in config.get("mappings", {}).values():
        required_sheets.add(str(mapping.get("sheetNumber", 1)))

    sheet_map = {str(i + 1): sheet for i, sheet in enumerate(wb.sheetnames)}

    for sheet_number in required_sheets:
        if sheet_number not in sheet_map:
            errors.append(f"❌ Sheet number {sheet_number} is missing in the input file.")
            continue
        
        sheet_name = sheet_map[sheet_number]
        ws = wb[sheet_name]
        rows = []

        for row in ws.iter_rows():
            formatted_row = []
            for cell in row:
                value = cell.value
                number_format = cell.number_format  # Get cell format
                formatted_row.append(format_cell_value(value, number_format))
            
            if any(formatted_row):  # Ignore empty rows
                rows.append(formatted_row)

        if len(rows) < 2:
            errors.append(f"❌ Sheet '{sheet_name}' (#{sheet_number}) must contain at least two non-empty rows.")
            continue

        headers = rows[0]
        sheet_data = []
        for row in rows[1:]:
            row_dict = {headers[i]: row[i] if i < len(row) else "" for i in range(len(headers))}
            sheet_data.append(row_dict)
        data[sheet_number] = sheet_data

    return data, errors

def read_config(config_path):
    """ Reads the JSON config file containing template paths and mappings. """
    if not os.path.exists(config_path):
        return {}, [f"❌ Config file '{config_path}' not found."]

    with open(config_path, 'r', encoding='utf-8') as file:
        config = json.load(file)

    required_keys = ["templatePath", "outputPath", "inputPath", "mappings"]
    missing_keys = [key for key in required_keys if key not in config]
    
    if missing_keys:
        return {}, [f"❌ Config file is missing required keys: {', '.join(missing_keys)}"]

    # Check for duplicate placeholders
    seen_placeholders = set()
    duplicates = []
    for placeholder in config["mappings"]:
        if placeholder in seen_placeholders:
            duplicates.append(placeholder)
        seen_placeholders.add(placeholder)

    if duplicates:
        return {}, [f"❌ Config contains duplicate placeholders: {', '.join(duplicates)}"]

    return config, []

def validate_docx_file(file_path):
    """ Validates if the given .docx file is readable and properly formatted. """
    if not os.path.exists(file_path):
        return False, f"❌ Template file '{file_path}' not found."

    try:
        docx.Document(file_path)
        return True, ""
    except Exception as e:
        return False, f"❌ Failed to read '{file_path}'. The file might be corrupted or not a valid .docx.\n{e}"

def replace_text_preserving_format(paragraph, replacements, placeholder_counts, missing_placeholders, doc):
    """ Replaces placeholders in a paragraph while preserving formatting. """
    for placeholder, replacement in replacements.items():
        if replacement["type"] == "string":
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement["value"])
                    placeholder_counts[placeholder] += 1
                    missing_placeholders.discard(placeholder)
        elif replacement["type"] == "table":
            if placeholder in paragraph.text:
                parts = paragraph.text.split(placeholder)
                if parts[0]:
                    paragraph.text = parts[0]
                else:
                    paragraph.clear()
                            
                table_paragraph = paragraph.insert_paragraph_before()
                data = replacement["value"]
                if data:
                    num_rows = len(data) + 1
                    num_cols = len(data[0])
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                        
                    table_paragraph._element.addnext(table._element)
                    tbl = table._element
                    tblBorders = parse_xml(
                        '<w:tblBorders %s>'
                        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                        '</w:tblBorders>' % nsdecls('w')
                    )
                    tbl.tblPr.append(tblBorders)
                        
                    headers = data[0].keys()
                    for i, header in enumerate(headers):
                        cell = table.cell(0, i)
                        cell.text = header
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                        
                    for row_idx, entry in enumerate(data, start=1):
                        for col_idx, (key, value) in enumerate(entry.items()):
                            table.cell(row_idx, col_idx).text = str(value)
                    
                if parts[1]:
                    paragraph.insert_paragraph_before(parts[1])
            
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
                placeholder_counts[placeholder] += 1
                missing_placeholders.discard(placeholder)

def replace_placeholders(config, data, error_messages):
    template_path = os.path.abspath(config["templatePath"])
    output_path = os.path.abspath(config["outputPath"])

    valid, error_message = validate_docx_file(template_path)
    if not valid:
        error_messages.append(error_message)
        return

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    try:
        doc = docx.Document(template_path)
    except Exception as e:
        error_messages.append(f"❌ Unable to open template file '{template_path}'.\n{e}")
        return

    # Load custom functions
    replacements = {}
    for placeholder, mapping in config["mappings"].items():
        sheet_number = str(mapping.get("sheetNumber", 1))  # Default to sheet 1 if missing
        sheet_data = data.get(sheet_number, {})  # Get data for the specified sheet

        if "inputField" in mapping:
            input_field = mapping["inputField"]
            if sheet_data and input_field in sheet_data[0]:
                replacements[placeholder] = {
                    "type": "string",
                    "value": sheet_data[0][input_field]
                }
            else:
                error_messages.append(f"⚠️  Expected column '{input_field}' based on config but it was not found in the input data.")
        elif "customOperation" in mapping:
            custom_function_name = mapping["customOperation"]
            try:
                custom_function = getattr(customUtil, custom_function_name)
                if (mapping.get("type") == "table"):
                    params = mapping.get("params", {})
                    table_data = custom_function(
                        sheet_data, 
                        start_row=params.get("start_row", 1), 
                        end_row=params.get("end_row", None), 
                        start_col=params.get("start_col", 1), 
                        end_col=params.get("end_col", None)
                    )
                    replacements[placeholder] = {
                        "type": "table",
                        "value": table_data
                    }
                else:
                    replacements[placeholder] = {
                        "type": "string",
                        "value": str(custom_function(sheet_data))
                    }
            except AttributeError:
                error_messages.append(f"❌ Custom function '{custom_function_name}' not found in customUtil.py.")
            except Exception as e:
                error_messages.append(f"❌ Error executing custom function '{custom_function_name}': {e}")

        # Apply custom formatting if ayn after default formatting
        if "formatter" in mapping:
            try:
                formatterName = mapping["formatter"]
                custom_formatter = getattr(formatters, formatterName)
                oldValue = replacements[placeholder]["value"]
                formattedValue = custom_formatter(oldValue)
                replacements[placeholder]["value"] = formattedValue
            except AttributeError:
                print(f"Error: Formatter '{formatterName}' not found in formatters module.")
            except KeyError:
                print(f"Error: Placeholder '{placeholder}' not found in replacements dictionary.")
            except Exception as e:
                print(f"Unexpected error while applying formatter '{formatterName}': {e}")

    placeholder_counts = {placeholder: 0 for placeholder in replacements}
    missing_placeholders = set(replacements.keys())

    # Replace in paragraphs
    for  para in doc.paragraphs:
        replace_text_preserving_format(para, replacements, placeholder_counts, missing_placeholders, doc)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for  para in cell.paragraphs:
                    replace_text_preserving_format(para, replacements, placeholder_counts, missing_placeholders, doc)

    # Log mappings and save document
    log_mappings( config, placeholder_counts,  missing_placeholders, error_messages,  replacements)

    try:
        doc.save(output_path)
        print(f"✅ Document successfully saved at: {output_path}")
    except Exception as e:
        error_messages.append(f"❌ Failed to save document to '{output_path}'.\n{e}")

def log_mappings( config, placeholder_counts,  missing_placeholders, error_messages, replacements=None):
    """ Logs the mappings and placeholder statistics for debugging. """

    print("\n📌 **Placeholder Replacement Log:**")

    table_placeholder_count = 0

    # Calculate the maximum length for each column
    max_input_len = max(
        len(mapping.get("inputField", f"{mapping.get('customOperation', '')}"))
        for mapping in config["mappings"].values()
    )
    max_placeholder_len = max(len(placeholder) for placeholder in config["mappings"]) if config["mappings"] else 0
    max_value_len = max(
        len(str(replacements[placeholder]["value"])) 
        for placeholder in replacements 
        if replacements[placeholder]["type"] == "string"
        ) if replacements else 0
    
    max_count_len = max(len(str(count)) for count in placeholder_counts.values()) if placeholder_counts else 0

    # Set column widths
    col_width = max(max_input_len, len("INPUT"))
    placeholder_width = max(max_placeholder_len, len("PLACEHOLDER"))
    value_width = max(max_value_len, len("VALUE"))
    count_width = max(max_count_len, len("COUNT"))

    # Print header
    print(f"\n{'INPUT'.ljust(col_width)}  {'PLACEHOLDER'.ljust(placeholder_width)}  {'VALUE'.ljust(value_width)}  {'COUNT'.rjust(count_width)}")

    # Print each mapping
    for placeholder, mapping in config["mappings"].items():

        replacementDefaultValue = {"type": "string", "value": ""}

        if placeholder in replacements:
            replacementDefaultValue = replacements[placeholder]

        if replacementDefaultValue["type"] == "table":
            table_placeholder_count += placeholder_counts[placeholder]
            continue

        if "inputField" in mapping:
            input_field = mapping["inputField"]
        elif "customOperation" in mapping:
            input_field = f"{mapping['customOperation']}"

        value = replacementDefaultValue.get("value", "ERROR")

        count = placeholder_counts.get(placeholder, 0)
        print(f"{input_field.ljust(col_width)}  {placeholder.ljust(placeholder_width)}  {str(value).ljust(value_width)}  {str(count).rjust(count_width)}")


    # Collect missing placeholders in the error log
    for placeholder in missing_placeholders:
        error_messages.append(f"❌ ERROR: Placeholder {placeholder} not found in the document.")

    total_changes = sum(placeholder_counts.values())
    # Log expected count mismatch if the field exists
    expected_count = config.get("expectedCount")
    if expected_count is not None and total_changes != expected_count:
        error_messages.append(f"❌ ERROR: Total placeholders replaced ({total_changes}) does not match expected count ({expected_count}).")

    # Log placeholders
    print(f"\n✅ Total placeholder values changed: {total_changes}.")
    if table_placeholder_count > 0:
        print(f"✅ Tables created {table_placeholder_count}.")


def shouldExit(forceExit = False):
    choice = input("\nPress ENTER to continue... (or 'q' to quit): ")
    if choice.lower() == 'q':
        print("\n🛑 Exiting...\n")
        return True
    
    return False | forceExit


def main():
    parser = argparse.ArgumentParser()

    # Argument Details
    # parser.add_argument("-v", "--verbose", action="store_true", help="Show detailed placeholder changes with line numbers and counts")
    # parser.add_argument("-lp","--loop", action="store_true", help="Keep running after completion, allowing re-selection")

    config_info = list_config_files(CONFIG_DIR)

    if not config_info:
        print("❌ No valid config files found. Exiting.")
        if shouldExit():
            return

    while True:
        config_file = get_user_selected_config(config_info, CONFIG_DIR)
        if not config_file:
            return  # Exit if the user chooses to quit

        config, config_errors = read_config(config_file)

        if config_errors:
            for error in config_errors:
                print(error)
            if shouldExit(True):
                return

        input_path = config.get("inputPath")
        if not input_path or not os.path.exists(input_path):
            print(f"❌ Invalid inputPath: {input_path}")
            if shouldExit(True):
                return
        
        data, data_errors = read_data(input_path,config)

        error_messages = data_errors + config_errors

        try:
            if data and config:
                replace_placeholders(config, data, error_messages)
        except Exception as e:
            print(f"Got internal error {e} {e.__traceback__}")


        if error_messages:
            print("\n📌 **Errors Encountered:** ")
            for error in error_messages:
                print(error)

        if shouldExit():
            return
        
        print("\n\n\n\n\n")
    
if __name__ == "__main__":
    main()