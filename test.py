import os
import docx
import json
import openpyxl
import csv
import argparse
import importlib.util
from collections import defaultdict

def list_config_files(config_dir="config"):
    """ Lists all available JSON config files and extracts product names. """
    if not os.path.exists(config_dir):
        print(f"❌ Config directory '{config_dir}' not found.")
        return []

    config_files = [f for f in os.listdir(config_dir) if f.endswith(".json")]
    config_info = []

    for file in config_files:
        file_path = os.path.join(config_dir, file)
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                config_data = json.load(f)
                product_name = config_data.get("productName", "Unknown Product")
                config_info.append((file, product_name))
        except Exception as e:
            print(f"⚠️  Skipping '{file}': Unable to read or invalid JSON format. ({e})")

    if not config_info:
        print(f"❌ No valid JSON config files found in '{config_dir}'.")
    
    return config_info

def get_user_selected_config(config_info, config_dir="config"):
    """ Prompts the user to select a config file from available options. """
    while True:
        print("\n📌 Available Configurations:")
        for idx, (file, product_name) in enumerate(config_info, start=1):
            print(f"{idx}. {product_name} ({file})")

        choice = input("\nEnter the number of the configuration to use (or 'q' to quit): ")
        if choice.lower() == 'q':
            print("\n🛑 Exiting selection...")
            return None
        try:
            choice = int(choice)
            if 1 <= choice <= len(config_info):
                return os.path.join(config_dir, config_info[choice - 1][0])
            else:
                print("❌ Invalid selection. Please enter a valid number.")
        except ValueError:
            print("❌ Invalid input. Please enter a number.")

def read_data(file_path):
    """ Reads input data from a CSV or XLSX file and converts it to a dictionary. """
    if not os.path.exists(file_path):
        return {}, [f"❌ Data file '{file_path}' not found."]
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == ".csv":
        return read_csv(file_path)
    elif file_ext in [".xls", ".xlsx"]:
        return read_xlsx(file_path)
    else:
        return {}, ["❌ Unsupported file format. Please use CSV or XLSX."]

def read_csv(file_path):
    """ Reads a CSV file and extracts data from the first two non-empty rows. """
    with open(file_path, newline='', encoding='utf-8') as file:
        reader = csv.reader(file)
        rows = [row for row in reader if any(cell.strip() for cell in row)]  # Remove empty rows
        
        if len(rows) < 2:
            return {}, ["❌ CSV file must contain at least two non-empty rows."]
        
        headers = [h.strip() for h in rows[0]]
        values = [v.strip() for v in rows[1]]
        
        return dict(zip(headers, values)), []

def read_xlsx(file_path):
    """ Reads an XLSX file and extracts data from the first two non-empty rows. """
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheet = wb.active
    
    rows = []
    for row in sheet.iter_rows(values_only=True):
        if any(cell for cell in row if cell is not None and str(cell).strip()):  # Remove empty rows
            rows.append([str(cell).strip() if cell is not None else "" for cell in row])
    
    if len(rows) < 2:
        return {}, ["❌ XLSX file must contain at least two non-empty rows."]
    
    headers = rows[0]
    values = rows[1]
    
    return dict(zip(headers, values)), []

def read_config(config_path):
    """ Reads the JSON config file containing template paths and mappings. """
    if not os.path.exists(config_path):
        return {}, [f"❌ Config file '{config_path}' not found."]

    with open(config_path, 'r', encoding='utf-8') as file:
        config = json.load(file)

    required_keys = ["templatePath", "outputPath", "inputPath", "mappings"]
    missing_keys = [key for key in required_keys if key not in config]
    
    if missing_keys:
        return {}, [f"❌ Config file is missing required keys: {', '.join(missing_keys)}"]

    # Check for duplicate placeholders
    seen_placeholders = set()
    duplicates = []
    for placeholder in config["mappings"]:
        if placeholder in seen_placeholders:
            duplicates.append(placeholder)
        seen_placeholders.add(placeholder)

    if duplicates:
        return {}, [f"❌ Config contains duplicate placeholders: {', '.join(duplicates)}"]

    return config, []

def validate_docx_file(file_path):
    """ Validates if the given .docx file is readable and properly formatted. """
    if not os.path.exists(file_path):
        return False, f"❌ Template file '{file_path}' not found."

    try:
        docx.Document(file_path)
        return True, ""
    except Exception as e:
        return False, f"❌ Failed to read '{file_path}'. The file might be corrupted or not a valid .docx.\n{e}"

def load_custom_util(custom_util_path):
    """ Dynamically loads the customUtil module from a given path. """
    spec = importlib.util.spec_from_file_location("customUtil", custom_util_path)
    custom_util = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(custom_util)
    return custom_util

def replace_text_preserving_format(paragraph, replacements, placeholder_counts, placeholder_line_counts, line_num, missing_placeholders):
    """ Replaces placeholders in a paragraph while preserving formatting. """
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)
                placeholder_counts[key] += 1
                placeholder_line_counts[key][line_num] = placeholder_line_counts[key].get(line_num, 0) + 1
                missing_placeholders.discard(key)

def replace_placeholders(config, data, error_messages, verbose=False):
    template_path = os.path.abspath(config["templatePath"])
    output_path = os.path.abspath(config["outputPath"])
    custom_util_path = os.path.abspath("customUtil.py")  # Adjust this path as needed

    valid, error_message = validate_docx_file(template_path)
    if not valid:
        error_messages.append(error_message)
        return

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    try:
        doc = docx.Document(template_path)
    except Exception as e:
        error_messages.append(f"❌ Unable to open template file '{template_path}'.\n{e}")
        return

    # Load custom functions
    custom_util = load_custom_util(custom_util_path)

    replacements = {}
    for placeholder, mapping in config["mappings"].items():
        if "inputField" in mapping:
            input_field = mapping["inputField"]
            if input_field in data:
                replacements[placeholder] = data[input_field]
            else:
                error_messages.append(f"⚠️  Expected column '{input_field}' based on config but it was not found in the input data.")
        elif "customOperation" in mapping:
            custom_function_name = mapping["customOperation"]
            try:
                custom_function = getattr(custom_util, custom_function_name)
                replacements[placeholder] = str(custom_function(data))
            except AttributeError:
                error_messages.append(f"❌ Custom function '{custom_function_name}' not found in customUtil.py.")
            except Exception as e:
                error_messages.append(f"❌ Error executing custom function '{custom_function_name}': {e}")

    placeholder_counts = {placeholder: 0 for placeholder in replacements}
    placeholder_line_counts = {placeholder: defaultdict(int) for placeholder in replacements}
    missing_placeholders = set(replacements.keys())

    # Replace in paragraphs
    for line_num, para in enumerate(doc.paragraphs, start=1):
        replace_text_preserving_format(para, replacements, placeholder_counts, placeholder_line_counts, line_num, missing_placeholders)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for line_num, para in enumerate(cell.paragraphs, start=1):
                    replace_text_preserving_format(para, replacements, placeholder_counts, placeholder_line_counts, line_num, missing_placeholders)

    # Log mappings and save document
    log_mappings(data, config, placeholder_counts, placeholder_line_counts, missing_placeholders, error_messages, verbose, replacements)

    try:
        doc.save(output_path)
        print(f"✅ Document successfully saved at: {output_path}")
    except Exception as e:
        error_messages.append(f"❌ Failed to save document to '{output_path}'.\n{e}")

def log_mappings(data, config, placeholder_counts, placeholder_line_counts, missing_placeholders, error_messages, verbose=False, replacements=None):
    """ Logs the mappings and placeholder statistics for debugging. """
    print("\n📌 **Placeholder Replacement Log:**")

    # Calculate the maximum length for each column
    max_input_len = max(
        len(mapping.get("inputField", f"{mapping.get('customOperation', '')}"))
        for mapping in config["mappings"].values()
    )
    max_placeholder_len = max(len(placeholder) for placeholder in config["mappings"]) if config["mappings"] else 0
    max_value_len = max(len(str(value)) for value in replacements.values()) if replacements else 0
    max_count_len = max(len(str(count)) for count in placeholder_counts.values()) if placeholder_counts else 0

    # Set column widths
    col_width = max(max_input_len, len("INPUT"))
    placeholder_width = max(max_placeholder_len, len("PLACEHOLDER"))
    value_width = max(max_value_len, len("VALUE"))
    count_width = max(max_count_len, len("COUNT"))

    # Print header
    print(f"\n{'INPUT'.ljust(col_width)}      {'PLACEHOLDER'.ljust(placeholder_width)}      {'VALUE'.ljust(value_width)}      {'COUNT'.rjust(count_width)}")

    # Print each mapping
    for placeholder, mapping in config["mappings"].items():
        if "inputField" in mapping:
            input_field = mapping["inputField"]
            value = data.get(input_field, "")
        elif "customOperation" in mapping:
            input_field = f"{mapping['customOperation']}"
            value = replacements.get(placeholder, "")

        count = placeholder_counts.get(placeholder, 0)
        print(f"{input_field.ljust(col_width)}  ->  {placeholder.ljust(placeholder_width)}  ->  {str(value).ljust(value_width)}  ->  {str(count).ljust(count_width)}")

        if verbose and placeholder in placeholder_line_counts:
            for line_num, line_count in placeholder_line_counts[placeholder].items():
                print(f"    ⮡  {line_count} Time{'s' if line_count > 1 else ''} at Line {line_num}")
            print("")

    # Collect missing placeholders in the error log
    for placeholder in missing_placeholders:
        error_messages.append(f"❌ ERROR: Placeholder {placeholder} not found in the document.")

    total_changes = sum(placeholder_counts.values())
    # Log expected count mismatch if the field exists
    expected_count = config.get("expectedCount")
    if expected_count is not None and total_changes != expected_count:
        error_messages.append(f"❌ ERROR: Total placeholders replaced ({total_changes}) does not match expected count ({expected_count}).")

    # Log placeholders
    print(f"\n✅ Total placeholder values changed: {total_changes}.")

def main():
    parser = argparse.ArgumentParser()

    # Argument Details
    parser.add_argument("-v", "--verbose", action="store_true", help="Show detailed placeholder changes with line numbers and counts")
    # parser.add_argument("-lp","--loop", action="store_true", help="Keep running after completion, allowing re-selection")
    
    args = parser.parse_args()

    config_dir = "config"
    config_info = list_config_files(config_dir)

    if not config_info:
        print("❌ No valid config files found. Exiting.")
        return

    while True:
        config_file = get_user_selected_config(config_info, config_dir)
        if not config_file:
            return  # Exit if the user chooses to quit

        config, config_errors = read_config(config_file)

        if config_errors:
            for error in config_errors:
                print(error)
            return

        input_path = config.get("inputPath")
        if not input_path or not os.path.exists(input_path):
            print(f"❌ Invalid inputPath: {input_path}")
            return
        
        data, data_errors = read_data(input_path)

        error_messages = data_errors + config_errors

        if data and config:
            replace_placeholders(config, data, error_messages, args.verbose)

        if error_messages:
            print("\n📌 **Errors Encountered:** ")
            for error in error_messages:
                print(error)

        choice = input("\nPress ENTER to continue... (or 'q' to quit): ")
        if choice.lower() == 'q':
            print("\n🛑 Exiting...\n")
            return
        
        print("\n\n\n\n\n")

        
if __name__ == "__main__":
    main()